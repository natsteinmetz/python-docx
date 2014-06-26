# encoding: utf-8

"""
Step implementations for paragraph-related features
"""

from behave import given, then, when

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.text import Paragraph

from helpers import saved_docx_path, test_text


TEST_STYLE = 'Heading1'


# given ===================================================

@given('a document containing three paragraphs')
def given_a_document_containing_three_paragraphs(context):
    document = Document()
    document.add_paragraph('foo')
    document.add_paragraph('bar')
    document.add_paragraph('baz')
    context.document = document


@given('a paragraph with content and formatting')
def given_a_paragraph_with_content_and_formatting(context):
    p_xml = """\
        <w:p %s>
          <w:pPr>
            <w:pStyle w:val="%s"/>
          </w:pPr>
          <w:r>
            <w:t>foobar</w:t>
          </w:r>
        </w:p>""" % (nsdecls('w'), TEST_STYLE)
    p = parse_xml(p_xml)
    context.paragraph = Paragraph(p)


# when ====================================================

@when('I add a run to the paragraph')
def when_add_new_run_to_paragraph(context):
    context.r = context.p.add_run()


@when('I add text to the run')
def when_I_add_text_to_the_run(context):
    context.r.add_text(test_text)


@when('I clear the paragraph content')
def when_I_clear_the_paragraph_content(context):
    context.paragraph.clear()


@when('I insert a paragraph above the second paragraph')
def when_I_insert_a_paragraph_above_the_second_paragraph(context):
    paragraph = context.document.paragraphs[1]
    paragraph.insert_paragraph_before('foobar', 'Heading1')


@when('I set the paragraph style')
def when_I_set_the_paragraph_style(context):
    context.paragraph.add_run().add_text(test_text)
    context.paragraph.style = TEST_STYLE


# then =====================================================

@then('the document contains four paragraphs')
def then_the_document_contains_four_paragraphs(context):
    assert len(context.document.paragraphs) == 4


@then('the document contains the text I added')
def then_document_contains_text_I_added(context):
    document = Document(saved_docx_path)
    paragraphs = document.paragraphs
    p = paragraphs[-1]
    r = p.runs[0]
    assert r.text == test_text


@then('the paragraph formatting is preserved')
def then_the_paragraph_formatting_is_preserved(context):
    assert context.paragraph.style == TEST_STYLE


@then('the paragraph has no content')
def then_the_paragraph_has_no_content(context):
    assert context.paragraph.text == ''


@then('the paragraph has the style I set')
def then_the_paragraph_has_the_style_I_set(context):
    paragraph = Document(saved_docx_path).paragraphs[-1]
    assert paragraph.style == TEST_STYLE


@then('the style of the second paragraph matches the style I set')
def then_the_style_of_the_second_paragraph_matches_the_style_I_set(context):
    second_paragraph = context.document.paragraphs[1]
    assert second_paragraph.style == 'Heading1'


@then('the text of the second paragraph matches the text I set')
def then_the_text_of_the_second_paragraph_matches_the_text_I_set(context):
    second_paragraph = context.document.paragraphs[1]
    assert second_paragraph.text == 'foobar'
